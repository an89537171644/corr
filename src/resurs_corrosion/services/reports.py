from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, List, Optional, Sequence, Tuple
from xml.sax.saxutils import escape

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from .. import __version__
from ..domain import (
    AnalysisRunRead,
    AxialForceKind,
    BaselineReportRequest,
    BaselineStoredElementRequest,
    CalculationRequest,
    CalculationResponse,
    CheckType,
    EngineeringCapacityMode,
    EngineeringConfidenceLevel,
    ForecastMode,
    NormativeCompletenessLevel,
    RateFitMode,
    ReducerMode,
    ReportArtifact,
    ReportBundle,
    ReportFormat,
    ResistanceMode,
    SectionType,
)
from ..models import ElementModel
from ..storage import build_calculation_request
from .engine import run_calculation


DEFAULT_REPORTS_DIR = Path("generated_reports")
MAX_TIMELINE_ROWS = 24


@dataclass
class ReportContext:
    title: str
    author: Optional[str]
    generated_at: datetime
    analysis_id: Optional[int]
    asset_id: Optional[int]
    element_db_id: Optional[int]
    calculation_request: CalculationRequest
    calculation_response: CalculationResponse


def generate_baseline_report_bundle(
    element: ElementModel,
    payload: BaselineReportRequest,
    reports_dir: Path,
    calculation_request: Optional[CalculationRequest] = None,
    calculation_response: Optional[CalculationResponse] = None,
    analysis_id: Optional[int] = None,
) -> ReportBundle:
    if calculation_request is None:
        calculation_request = build_calculation_request(
            element,
            BaselineStoredElementRequest(
                forecast_horizon_years=payload.forecast_horizon_years,
                time_step_years=payload.time_step_years,
                current_service_life_years=payload.current_service_life_years,
                scenarios=payload.scenarios,
                forecast_mode=payload.forecast_mode,
            ),
        )
    if calculation_response is None:
        calculation_response = run_calculation(calculation_request)

    generated_at = datetime.now(timezone.utc)
    context = ReportContext(
        title=payload.report_title or default_report_title(element.element_code),
        author=payload.author,
        generated_at=generated_at,
        analysis_id=analysis_id,
        asset_id=element.asset.id if element.asset is not None else None,
        element_db_id=element.id,
        calculation_request=calculation_request,
        calculation_response=calculation_response,
    )
    return generate_report_bundle(context, payload.output_formats, reports_dir)


def build_report_context_from_analysis(
    analysis_run: AnalysisRunRead,
    report_title: Optional[str] = None,
    author: Optional[str] = None,
) -> ReportContext:
    return ReportContext(
        title=report_title or default_report_title(analysis_run.request.element.element_id),
        author=author,
        generated_at=analysis_run.generated_at,
        analysis_id=analysis_run.id,
        asset_id=analysis_run.asset_id,
        element_db_id=analysis_run.element_id,
        calculation_request=analysis_run.request,
        calculation_response=analysis_run.result,
    )


def generate_report_bundle(
    context: ReportContext,
    output_formats: Sequence[ReportFormat],
    reports_dir: Path,
) -> ReportBundle:
    target_dir = build_report_storage_dir(reports_dir, context)
    target_dir.mkdir(parents=True, exist_ok=True)
    stem = build_report_stem(context.title, context.element_db_id or context.analysis_id or 0, context.generated_at)

    artifacts: List[ReportArtifact] = []
    for report_format in output_formats:
        file_path = target_dir / f"{stem}.{report_format.value}"
        write_report_file(context, report_format, file_path)
        artifacts.append(
            ReportArtifact(
                format=report_format,
                filename=file_path.name,
                media_type=media_type_for_report_format(report_format),
                size_bytes=file_path.stat().st_size,
                file_path=str(file_path.resolve()),
                download_url=build_download_url(context, file_path.name),
            )
        )

    return ReportBundle(
        analysis_id=context.analysis_id,
        report_title=context.title,
        generated_at=context.generated_at,
        environment_category=context.calculation_response.environment_category,
        forecast_mode=context.calculation_response.forecast_mode,
        scenario_count=len(context.calculation_response.results),
        recommended_action=context.calculation_response.risk_profile.recommended_action,
        artifacts=artifacts,
    )


def build_report_storage_dir(reports_dir: Path, context: ReportContext) -> Path:
    if context.element_db_id is not None:
        return reports_dir / f"element-{context.element_db_id}"
    if context.analysis_id is not None:
        return reports_dir / f"analysis-{context.analysis_id}"
    return reports_dir / "adhoc"


def build_download_url(context: ReportContext, filename: str) -> str:
    if context.element_db_id is not None:
        return f"/api/v1/reports/{context.element_db_id}/{filename}"
    return ""


def media_type_for_report_format(report_format: ReportFormat) -> str:
    if report_format == ReportFormat.DOCX:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if report_format == ReportFormat.PDF:
        return "application/pdf"
    if report_format == ReportFormat.HTML:
        return "text/html"
    return "text/markdown"


def write_report_file(context: ReportContext, report_format: ReportFormat, file_path: Path) -> None:
    if report_format == ReportFormat.DOCX:
        write_docx_report(context, file_path)
    elif report_format == ReportFormat.PDF:
        write_pdf_report(context, file_path)
    elif report_format == ReportFormat.HTML:
        write_html_report(context, file_path)
    else:
        write_markdown_report(context, file_path)


def default_report_title(element_code: str) -> str:
    return f"Отчет по остаточному ресурсу для {element_code}"


def build_report_stem(report_title: str, reference_id: int, generated_at: datetime) -> str:
    normalized = unicodedata.normalize("NFKD", report_title).encode("ascii", "ignore").decode("ascii")
    slug = re.sub(r"[^A-Za-z0-9]+", "-", normalized).strip("-").lower()
    if not slug:
        slug = "otchet-po-ostatochnomu-resursu"
    timestamp = generated_at.strftime("%Y%m%d-%H%M%S")
    return f"{slug}-{reference_id}-{timestamp}"


def write_docx_report(context: ReportContext, file_path: Path) -> None:
    document = Document()
    document.add_heading(context.title, 0)
    document.add_paragraph(f"Сформирован: {format_datetime(context.generated_at)}")
    if context.author:
        document.add_paragraph(f"Подготовил: {context.author}")

    add_docx_paragraph(document, build_scope_line(context))
    add_docx_paragraph(document, build_model_line(context))

    document.add_heading("Идентификация объекта и элемента", level=1)
    add_docx_table(document, ["Поле", "Значение"], build_identification_rows(context))

    document.add_heading("Исходные данные и наблюдения", level=1)
    add_docx_table(document, ["Поле", "Значение"], build_input_rows(context))
    add_docx_table(document, ["Параметр обследования", "Значение"], build_inspection_rows(context))
    add_docx_table(
        document,
        ["Зона", "Точка", "Толщина, мм", "Погрешность, мм", "Качество", "Комментарий"],
        build_measurement_rows(context),
    )

    document.add_heading("Модель и коэффициенты", level=1)
    add_docx_table(document, ["Параметр", "Значение"], build_model_rows(context))

    document.add_heading("Текущее состояние зон", level=1)
    add_docx_table(
        document,
        ["Зона", "Роль", "Наблюдаемая потеря, мм", "Прогнозная потеря, мм", "Скорость прогноза, мм/год", "Режим", "Эффективная толщина, мм"],
        build_zone_rows(context),
    )
    document.add_heading("Диагностика скорости деградации", level=1)
    add_docx_table(
        document,
        ["Зона", "Источник", "Режим fit", "n", "Сумма весов", "RMSE", "R2-like", "История, лет", "Интервал v"],
        build_rate_diagnostic_rows(context),
    )

    document.add_heading("Сравнение сценариев", level=1)
    add_docx_table(
        document,
        ["Сценарий", "Несущая способность", "Воздействие", "Запас", "Остаточный ресурс, лет", "Предельное состояние"],
        build_scenario_rows(context),
    )
    document.add_heading("Диагностика остаточного ресурса", level=1)
    add_docx_table(
        document,
        ["Сценарий", "Номинал, лет", "Консервативно, лет", "Верхняя граница, лет", "Интервал", "Статус", "Поиск", "Ширина интервала", "Итерации"],
        build_scenario_detail_rows(context),
    )
    document.add_heading("Неопределенность и риск", level=1)
    add_docx_table(document, ["Параметр", "Значение"], build_uncertainty_rows(context))
    document.add_heading("Матрица предупреждений", level=1)
    add_docx_table(document, ["Серьезность", "Источник", "Содержание"], build_warning_group_rows(context))

    document.add_heading("Фрагмент временной диаграммы", level=1)
    add_docx_table(
        document,
        ["Возраст, лет", "R central", "R conservative", "R upper", "S", "Запас"],
        build_timeline_rows(context),
    )

    document.add_heading("Ограничения применимости", level=1)
    for line in build_limitation_lines(context):
        add_docx_paragraph(document, line)

    document.add_heading("Рекомендации", level=1)
    for line in build_recommendation_lines(context):
        add_docx_paragraph(document, line)

    document.save(file_path)


def write_pdf_report(context: ReportContext, file_path: Path) -> None:
    regular_font, bold_font = register_pdf_fonts()
    styles = get_pdf_styles(regular_font, bold_font)
    story = []

    story.append(Paragraph(escape(context.title), styles["title"]))
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph(escape(f"Сформирован: {format_datetime(context.generated_at)}"), styles["body"]))
    if context.author:
        story.append(Paragraph(escape(f"Подготовил: {context.author}"), styles["body"]))
    story.append(Paragraph(escape(build_scope_line(context)), styles["body"]))
    story.append(Paragraph(escape(build_model_line(context)), styles["body"]))
    story.append(Spacer(1, 4 * mm))

    story.extend(build_pdf_section("Идентификация объекта и элемента", ["Поле", "Значение"], build_identification_rows(context), styles))
    story.extend(build_pdf_section("Исходные данные и наблюдения", ["Поле", "Значение"], build_input_rows(context), styles))
    story.extend(build_pdf_section("Сводка по обследованию", ["Параметр обследования", "Значение"], build_inspection_rows(context), styles))
    story.extend(
        build_pdf_section(
            "Замеры",
            ["Зона", "Точка", "Толщина, мм", "Погрешность, мм", "Качество", "Комментарий"],
            build_measurement_rows(context),
            styles,
        )
    )
    story.extend(build_pdf_section("Модель и коэффициенты", ["Параметр", "Значение"], build_model_rows(context), styles))
    story.extend(
        build_pdf_section(
            "Текущее состояние зон",
            ["Зона", "Роль", "Наблюдаемая потеря, мм", "Прогнозная потеря, мм", "Скорость прогноза, мм/год", "Режим", "Эффективная толщина, мм"],
            build_zone_rows(context),
            styles,
        )
    )
    story.extend(
        build_pdf_section(
            "Диагностика скорости деградации",
            ["Зона", "Источник", "Режим fit", "n", "Сумма весов", "RMSE", "R2-like", "История, лет", "Интервал v"],
            build_rate_diagnostic_rows(context),
            styles,
        )
    )
    story.extend(
        build_pdf_section(
            "Сравнение сценариев",
            ["Сценарий", "Несущая способность", "Воздействие", "Запас", "Остаточный ресурс, лет", "Предельное состояние"],
            build_scenario_rows(context),
            styles,
        )
    )
    story.extend(
        build_pdf_section(
            "Диагностика остаточного ресурса",
            ["Сценарий", "Номинал, лет", "Консервативно, лет", "Верхняя граница, лет", "Интервал", "Статус", "Поиск", "Ширина интервала", "Итерации"],
            build_scenario_detail_rows(context),
            styles,
        )
    )
    story.extend(build_pdf_section("Неопределенность и риск", ["Параметр", "Значение"], build_uncertainty_rows(context), styles))
    story.extend(
        build_pdf_section(
            "Матрица предупреждений",
            ["Серьезность", "Источник", "Содержание"],
            build_warning_group_rows(context),
            styles,
        )
    )
    story.extend(
        build_pdf_section(
            "Фрагмент временной диаграммы",
            ["Возраст, лет", "R central", "R conservative", "R upper", "S", "Запас"],
            build_timeline_rows(context),
            styles,
        )
    )
    story.extend(build_pdf_list_section("Ограничения применимости", build_limitation_lines(context), styles))

    story.append(Paragraph("Рекомендации", styles["heading"]))
    for line in build_recommendation_lines(context):
        story.append(Paragraph(escape(line), styles["body"]))

    doc = SimpleDocTemplate(
        str(file_path),
        pagesize=A4,
        leftMargin=15 * mm,
        rightMargin=15 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )
    doc.build(story)


def write_markdown_report(context: ReportContext, file_path: Path) -> None:
    file_path.write_text(build_markdown_report(context), encoding="utf-8")


def write_html_report(context: ReportContext, file_path: Path) -> None:
    file_path.write_text(build_html_report(context), encoding="utf-8")


def build_markdown_report(context: ReportContext) -> str:
    lines = [
        f"# {context.title}",
        "",
        f"- Сформирован: {format_datetime(context.generated_at)}",
    ]
    if context.author:
        lines.append(f"- Подготовил: {context.author}")
    lines.extend(
        [
            f"- Область оценки: {build_scope_line(context)}",
            f"- Расчетная схема: {build_model_line(context)}",
            "",
            "## Идентификация объекта и элемента",
            "",
            render_markdown_table(["Поле", "Значение"], build_identification_rows(context)),
            "",
            "## Исходные данные и наблюдения",
            "",
            render_markdown_table(["Поле", "Значение"], build_input_rows(context)),
            "",
            render_markdown_table(["Параметр обследования", "Значение"], build_inspection_rows(context)),
            "",
            render_markdown_table(
                ["Зона", "Точка", "Толщина, мм", "Погрешность, мм", "Качество", "Комментарий"],
                build_measurement_rows(context),
            ),
            "",
            "## Модель и коэффициенты",
            "",
            render_markdown_table(["Параметр", "Значение"], build_model_rows(context)),
            "",
            "## Текущее состояние зон",
            "",
            render_markdown_table(
                ["Зона", "Роль", "Наблюдаемая потеря, мм", "Прогнозная потеря, мм", "Скорость прогноза, мм/год", "Режим", "Эффективная толщина, мм"],
                build_zone_rows(context),
            ),
            "",
            "## Диагностика скорости деградации",
            "",
            render_markdown_table(
                ["Зона", "Источник", "Режим fit", "n", "Сумма весов", "RMSE", "R2-like", "История, лет", "Интервал v"],
                build_rate_diagnostic_rows(context),
            ),
            "",
            "## Сравнение сценариев",
            "",
            render_markdown_table(
                ["Сценарий", "Несущая способность", "Воздействие", "Запас", "Остаточный ресурс, лет", "Предельное состояние"],
                build_scenario_rows(context),
            ),
            "",
            "## Диагностика остаточного ресурса",
            "",
            render_markdown_table(
                ["Сценарий", "Номинал, лет", "Консервативно, лет", "Верхняя граница, лет", "Интервал", "Статус", "Поиск", "Ширина интервала", "Итерации"],
                build_scenario_detail_rows(context),
            ),
            "",
            "## Неопределенность и риск",
            "",
            render_markdown_table(["Параметр", "Значение"], build_uncertainty_rows(context)),
            "",
            "## Матрица предупреждений",
            "",
            render_markdown_table(["Серьезность", "Источник", "Содержание"], build_warning_group_rows(context)),
            "",
            "## Фрагмент временной диаграммы",
            "",
            render_markdown_table(["Возраст, лет", "R central", "R conservative", "R upper", "S", "Запас"], build_timeline_rows(context)),
            "",
            "## Ограничения применимости",
            "",
        ]
    )
    lines.extend([f"- {line}" for line in build_limitation_lines(context)])
    lines.extend(
        [
            "",
            "## Рекомендации",
            "",
        ]
    )
    lines.extend([f"- {line}" for line in build_recommendation_lines(context)])
    lines.append("")
    return "\n".join(lines)


def build_html_report(context: ReportContext) -> str:
    sections = [
        ("Идентификация объекта и элемента", ["Поле", "Значение"], build_identification_rows(context)),
        ("Исходные данные и наблюдения", ["Поле", "Значение"], build_input_rows(context)),
        ("Сводка по обследованию", ["Параметр обследования", "Значение"], build_inspection_rows(context)),
        (
            "Замеры",
            ["Зона", "Точка", "Толщина, мм", "Погрешность, мм", "Качество", "Комментарий"],
            build_measurement_rows(context),
        ),
        ("Модель и коэффициенты", ["Параметр", "Значение"], build_model_rows(context)),
        (
            "Текущее состояние зон",
            ["Зона", "Роль", "Наблюдаемая потеря, мм", "Прогнозная потеря, мм", "Скорость прогноза, мм/год", "Режим", "Эффективная толщина, мм"],
            build_zone_rows(context),
        ),
        (
            "Диагностика скорости деградации",
            ["Зона", "Источник", "Режим fit", "n", "Сумма весов", "RMSE", "R2-like", "История, лет", "Интервал v"],
            build_rate_diagnostic_rows(context),
        ),
        (
            "Сравнение сценариев",
            ["Сценарий", "Несущая способность", "Воздействие", "Запас", "Остаточный ресурс, лет", "Предельное состояние"],
            build_scenario_rows(context),
        ),
        (
            "Диагностика остаточного ресурса",
            ["Сценарий", "Номинал, лет", "Консервативно, лет", "Верхняя граница, лет", "Интервал", "Статус", "Поиск", "Ширина интервала", "Итерации"],
            build_scenario_detail_rows(context),
        ),
        ("Неопределенность и риск", ["Параметр", "Значение"], build_uncertainty_rows(context)),
        ("Матрица предупреждений", ["Серьезность", "Источник", "Содержание"], build_warning_group_rows(context)),
        ("Фрагмент временной диаграммы", ["Возраст, лет", "R central", "R conservative", "R upper", "S", "Запас"], build_timeline_rows(context)),
    ]

    body = [
        "<!doctype html>",
        "<html lang=\"ru\">",
        "<head>",
        "  <meta charset=\"utf-8\">",
        f"  <title>{escape(context.title)}</title>",
        "  <style>",
        "    body { font-family: 'Segoe UI', Arial, sans-serif; margin: 32px; color: #18212f; background: #f4f7fb; }",
        "    main { max-width: 1080px; margin: 0 auto; background: #ffffff; padding: 32px; border-radius: 18px; box-shadow: 0 18px 48px rgba(24, 33, 47, 0.12); }",
        "    h1, h2 { color: #102544; }",
        "    p.meta { margin: 0 0 6px; color: #51627a; }",
        "    table { width: 100%; border-collapse: collapse; margin: 14px 0 24px; font-size: 14px; }",
        "    th, td { border: 1px solid #c6d2e1; padding: 8px 10px; vertical-align: top; text-align: left; }",
        "    th { background: #dce7f5; }",
        "    ul { padding-left: 22px; }",
        "  </style>",
        "</head>",
        "<body>",
        "  <main>",
        f"    <h1>{escape(context.title)}</h1>",
        f"    <p class=\"meta\">Сформирован: {escape(format_datetime(context.generated_at))}</p>",
    ]
    if context.author:
        body.append(f"    <p class=\"meta\">Подготовил: {escape(context.author)}</p>")
    body.extend(
        [
            f"    <p class=\"meta\">Область оценки: {escape(build_scope_line(context))}</p>",
            f"    <p class=\"meta\">Расчетная схема: {escape(build_model_line(context))}</p>",
        ]
    )
    for title, headers, rows in sections:
        body.append(f"    <h2>{escape(title)}</h2>")
        body.append(render_html_table(headers, rows))
    body.append("    <h2>Ограничения применимости</h2>")
    body.append("    <ul>")
    body.extend([f"      <li>{escape(line)}</li>" for line in build_limitation_lines(context)])
    body.append("    </ul>")
    body.append("    <h2>Рекомендации</h2>")
    body.append("    <ul>")
    body.extend([f"      <li>{escape(line)}</li>" for line in build_recommendation_lines(context)])
    body.extend(
        [
            "    </ul>",
            "  </main>",
            "</body>",
            "</html>",
        ]
    )
    return "\n".join(body)


def build_scope_line(context: ReportContext) -> str:
    request = context.calculation_request
    return (
        f"Объект: {request.asset.name}; элемент: {request.element.element_id}; "
        f"тип: {request.element.element_type}; среда: {request.environment_category.value}."
    )


def build_model_line(context: ReportContext) -> str:
    return (
        "Отчет сформирован по инженерной цепочке "
        "delta_obs -> v_z -> прогноз потери -> t_eff -> эффективное сечение -> несущая способность -> остаточный ресурс, "
        "при сохранении классического атмосферного закона как базовой и резервной модели."
    )


def build_identification_rows(context: ReportContext) -> List[List[str]]:
    request = context.calculation_request
    return [
        ["ID объекта", str(context.asset_id or "-")],
        ["Наименование объекта", request.asset.name],
        ["Адрес", request.asset.address or "-"],
        ["Год ввода в эксплуатацию", str(request.asset.commissioned_year or "-")],
        ["Назначение", request.asset.purpose or "-"],
        ["Класс ответственности", request.asset.responsibility_class or "-"],
        ["ID элемента в БД", str(context.element_db_id or "-")],
        ["Код элемента", request.element.element_id],
        ["Тип элемента", request.element.element_type],
        ["Марка стали", request.element.steel_grade or "-"],
        ["Расчетная схема", request.element.work_scheme or "-"],
        ["Зона эксплуатации", request.element.operating_zone or "-"],
        ["ID расчета", str(context.analysis_id or "-")],
    ]


def build_input_rows(context: ReportContext) -> List[List[str]]:
    request = context.calculation_request
    rows = [
        ["Категория среды", request.environment_category.value],
        ["Режим прогноза", translate_forecast_mode(context.calculation_response.forecast_mode)],
        ["Текущий срок службы, лет", format_number(request.current_service_life_years, 2)],
        ["Горизонт прогноза, лет", format_number(request.forecast_horizon_years, 2)],
        ["Шаг по времени, лет", format_number(request.time_step_years, 2)],
        ["Тип сечения", request.section.section_type.value],
        ["Предел текучести fy, МПа", format_number(request.material.fy_mpa, 2)],
        ["Gamma_m", format_number(request.material.gamma_m, 3)],
        ["Коэффициент устойчивости", format_number(request.material.stability_factor, 3)],
        ["Количество зон", str(len(request.zones))],
        ["Количество обследований", str(len(request.inspections))],
    ]
    rows[9:9] = build_action_rows(request)
    return rows


def build_inspection_rows(context: ReportContext) -> List[List[str]]:
    inspections = sorted(context.calculation_request.inspections, key=lambda item: item.performed_at, reverse=True)
    if not inspections:
        return [["Статус", "Обследования отсутствуют."]]

    latest = inspections[0]
    return [
        ["Количество обследований", str(len(inspections))],
        ["Код последнего обследования", latest.inspection_id],
        ["Дата последнего обследования", latest.performed_at.isoformat()],
        ["Метод последнего обследования", latest.method],
        ["Исполнитель", latest.executor or "-"],
        ["Основные выводы", latest.findings or "-"],
    ]


def build_measurement_rows(context: ReportContext) -> List[List[str]]:
    rows: List[List[str]] = []
    inspections = sorted(context.calculation_request.inspections, key=lambda item: item.performed_at, reverse=True)
    for inspection in inspections:
        for measurement in inspection.measurements:
            rows.append(
                [
                    measurement.zone_id,
                    measurement.point_id or "-",
                    format_number(measurement.thickness_mm, 3),
                    format_number(measurement.error_mm, 3),
                    format_number(measurement.quality, 2),
                    measurement.comment or "-",
                ]
            )

    if rows:
        return rows
    return [["-", "-", "-", "-", "-", "-"]]


def build_model_rows(context: ReportContext) -> List[List[str]]:
    response = context.calculation_response
    coefficients = response.environment_coefficients
    rows = [
        ["Коэффициент среды k, мм", format_number(coefficients["k_mm"], 4)],
        ["Показатель времени b", format_number(coefficients["b"], 4)],
        ["Консервативный показатель b", format_number(coefficients["b_conservative"], 4)],
        ["Режим прогноза", translate_forecast_mode(response.forecast_mode)],
        ["Класс инженерной уверенности", translate_confidence_level(response.engineering_confidence_level)],
        ["Режим редьюсера", translate_reducer_mode(response.reducer_mode)],
        ["Режим сопротивления", translate_resistance_mode(response.resistance_mode)],
        ["Режим аппроксимации скорости", translate_rate_fit_mode(response.rate_fit_mode)],
        ["Режим ML-контура", translate_ml_mode(response.ml_mode)],
        ["Использовано обследований", str(response.used_inspection_count)],
        ["Использовано замеров", str(response.used_measurement_count)],
        ["ML-модель", response.ml_model_version.name],
        ["Версия ML-модели", response.ml_model_version.version],
        ["Примечание по ML-модели", response.ml_model_version.notes or "-"],
        ["ML execution mode", response.ml_model_version.execution_mode or "-"],
        ["ML blend mode", response.ml_model_version.blend_mode or "-"],
        ["ML interval source", response.ml_model_version.interval_source or "-"],
        ["ML accepted/rejected rows", f"{response.ml_model_version.accepted_row_count}/{response.ml_model_version.rejected_row_count}"],
        ["ML accepted/rejected candidates", f"{response.ml_model_version.accepted_candidate_count}/{response.ml_model_version.rejected_candidate_count}"],
        ["ML acceptance policy", format_acceptance_policy(response.ml_model_version.acceptance_policy)],
        ["Версия набора данных", response.dataset_version.code],
        ["Источник данных", response.dataset_version.source],
        ["Хэш набора данных", response.dataset_version.data_hash or "-"],
        ["Accepted/rejected rows dataset", f"{response.dataset_version.accepted_row_count}/{response.dataset_version.rejected_row_count}"],
        ["Количество сценариев", str(len(response.results))],
        ["Доля превышений", format_number(response.risk_profile.exceedance_share, 3)],
    ]
    rows.extend(
        [
            ["Engineering capacity mode", getattr(response.engineering_capacity_mode, "value", str(response.engineering_capacity_mode))],
            ["Normative completeness", getattr(response.normative_completeness_level, "value", str(response.normative_completeness_level))],
            ["ML correction factor", format_number(response.ml_correction_factor, 3)],
            ["ML coverage score", format_number(response.coverage_score, 3)],
            ["ML training regime", response.training_regime],
            ["Normalization mode", response.normalization_mode],
        ]
    )
    if response.ml_model_version.dataset_journal:
        rows.append(["ML dataset journal", summarize_dataset_journal(response.ml_model_version.dataset_journal)])
    if response.ml_model_version.candidate_registry:
        rows.append(["ML candidate registry", summarize_candidate_registry(response.ml_model_version.candidate_registry)])
    return rows


def build_zone_rows(context: ReportContext) -> List[List[str]]:
    baseline = context.calculation_response.results[0]
    return [
        [
            state.zone_id,
            state.role,
            format_optional_number(state.observed_loss_mm, 3),
            format_number(state.corrosion_loss_mm, 3),
            format_optional_number(state.forecast_rate_mm_per_year, 4),
            state.forecast_source,
            format_number(state.effective_thickness_mm, 3),
        ]
        for state in baseline.zone_states
    ]


def build_scenario_rows(context: ReportContext) -> List[List[str]]:
    rows: List[List[str]] = []
    for result in context.calculation_response.results:
        rows.append(
            [
                result.scenario_name,
                f"{format_number(result.resistance_value, 3)} {result.resistance_unit}",
                f"{format_number(result.demand_value, 3)} {result.demand_unit}",
                format_number(result.margin_value, 3),
                format_optional_number(result.remaining_life_years, 2),
                "Достигнуто" if result.limit_state_reached_within_horizon else "Не достигнуто",
            ]
        )
    return rows


def build_rate_diagnostic_rows(context: ReportContext) -> List[List[str]]:
    rows: List[List[str]] = []
    for observation in context.calculation_response.zone_observations:
        rate_band = "-"
        if observation.rate_lower_mm_per_year is not None or observation.rate_upper_mm_per_year is not None:
            rate_band = (
                f"{format_optional_number(observation.rate_lower_mm_per_year, 4)} .. "
                f"{format_optional_number(observation.rate_upper_mm_per_year, 4)}"
            )
        rows.append(
            [
                observation.zone_id,
                observation.source,
                translate_rate_fit_mode(observation.rate_fit_mode),
                str(observation.fit_sample_size or observation.used_points_count),
                format_optional_number(observation.effective_weight_sum, 2),
                format_optional_number(observation.fit_rmse, 4),
                format_optional_number(observation.fit_r2_like, 3),
                format_optional_number(observation.history_span_years, 2),
                rate_band,
            ]
        )
    return rows or [["-", "-", "-", "-", "-", "-", "-", "-", "-"]]


def build_scenario_detail_rows(context: ReportContext) -> List[List[str]]:
    rows: List[List[str]] = []
    for result in context.calculation_response.results:
        interval = format_life_interval(result.life_interval_years)
        rows.append(
            [
                result.scenario_name,
                format_optional_number(result.remaining_life_nominal_years, 2),
                format_optional_number(result.remaining_life_conservative_years, 2),
                format_optional_number(result.life_interval_years.upper_years, 2),
                interval,
                translate_refinement_status(result.refinement_diagnostics.status),
                translate_crossing_mode(result.crossing_search_mode),
                format_optional_number(result.crossing_bracket_width_years, 2),
                str(result.crossing_refinement_iterations),
            ]
        )
    return rows or [["-", "-", "-", "-", "-", "-", "-", "-", "-"]]


def build_uncertainty_rows(context: ReportContext) -> List[List[str]]:
    response = context.calculation_response
    return [
        ["Режим риска", translate_risk_mode(response.risk_mode)],
        ["Интервал остаточного ресурса", format_life_interval(response.life_interval_years)],
        ["Номинальный ресурс, лет", format_optional_number(response.life_interval_years.nominal_years, 2)],
        ["Консервативный ресурс, лет", format_optional_number(response.life_interval_years.conservative_years, 2)],
        ["Верхняя граница ресурса, лет", format_optional_number(response.life_interval_years.upper_years, 2)],
        ["Уровень uncertainty", translate_uncertainty_level(response.uncertainty_level)],
        ["Источник uncertainty", translate_uncertainty_source(response.uncertainty_source)],
        ["Режим поиска предельного состояния", translate_crossing_mode(response.crossing_search_mode)],
        ["Статус уточнения", translate_refinement_status(response.refinement_diagnostics.status)],
        ["Ширина интервала корня, лет", format_optional_number(response.refinement_diagnostics.bracket_width_years, 2)],
        ["Размах margin(t)", format_optional_number(response.refinement_diagnostics.margin_span_value, 3)],
        ["ML candidates", str(response.ml_candidate_count)],
        ["ML blend mode", response.ml_blend_mode],
        ["ML interval source", response.ml_interval_source],
        ["Траектории guidance", build_trajectory_summary(response.governing_uncertainty_trajectories)],
        ["Основания uncertainty band", "; ".join(response.uncertainty_basis) or "-"],
        ["Предупреждения uncertainty band", "; ".join(response.uncertainty_warnings) or "-"],
        ["Предупреждения refinement", "; ".join(response.refinement_diagnostics.warnings) or "-"],
    ]


def build_warning_group_rows(context: ReportContext) -> List[List[str]]:
    groups = collect_warning_groups(context.calculation_response)
    if not groups:
        return [["low", "none", "Дополнительные предупреждения не зарегистрированы."]]

    rows: List[List[str]] = []
    for severity, origin, items in groups:
        rows.append([severity, origin, "; ".join(items)])
    return rows


def collect_warning_groups(response: CalculationResponse) -> List[Tuple[str, str, List[str]]]:
    groups: List[Tuple[str, str, List[str]]] = []

    append_warning_group(groups, "medium", "calculation", response.warnings)
    append_warning_group(
        groups,
        "high" if response.fallback_flags else "medium",
        "fallback",
        [translate_fallback_flag(flag) for flag in response.fallback_flags],
    )
    append_warning_group(groups, "medium", "uncertainty", response.uncertainty_warnings)
    append_warning_group(groups, "medium", "refinement", response.refinement_diagnostics.warnings)

    for observation in response.zone_observations:
        append_warning_group(groups, "medium", f"rate_fit:{observation.zone_id}", observation.warnings)

    return groups


def append_warning_group(
    groups: List[Tuple[str, str, List[str]]],
    severity: str,
    origin: str,
    items: Sequence[str],
) -> None:
    normalized_items = [str(item) for item in items if str(item).strip()]
    if not normalized_items:
        return
    groups.append((severity, origin, unique_lines(normalized_items)))


def build_timeline_rows(context: ReportContext) -> List[List[str]]:
    baseline = context.calculation_response.results[0]
    central = baseline.uncertainty_trajectories.central or baseline.timeline
    conservative_map = {
        round(point.age_years, 6): point for point in (baseline.uncertainty_trajectories.conservative or baseline.timeline)
    }
    upper_map = {
        round(point.age_years, 6): point for point in (baseline.uncertainty_trajectories.upper or baseline.timeline)
    }
    sampled_timeline = sample_timeline(central, MAX_TIMELINE_ROWS)
    rows: List[List[str]] = []
    for point in sampled_timeline:
        key = round(point.age_years, 6)
        conservative_point = conservative_map.get(key, point)
        upper_point = upper_map.get(key, point)
        rows.append(
            [
                format_number(point.age_years, 2),
                format_number(point.resistance_value, 3),
                format_number(conservative_point.resistance_value, 3),
                format_number(upper_point.resistance_value, 3),
                format_number(point.demand_value, 3),
                format_number(point.margin_value, 3),
            ]
        )
    return rows


def build_recommendation_lines(context: ReportContext) -> List[str]:
    return [
        context.calculation_response.risk_profile.recommended_action,
        f"Рекомендуемый срок до следующего обследования: {format_number(context.calculation_response.risk_profile.next_inspection_within_years, 2)} лет.",
        context.calculation_response.risk_profile.method_note,
        f"Версия расчетного ядра: {__version__}",
    ]


def build_limitation_lines(context: ReportContext) -> List[str]:
    response = context.calculation_response
    lines: List[str] = []

    if response.reducer_mode == ReducerMode.GENERIC_FALLBACK:
        lines.append(
            "Для эффективного сечения использован режим generic_reduced. "
            "Это жестко маркированный fallback-режим по исходным пользовательским характеристикам, "
            "а не прямой reducer нормативного профиля."
        )

    if context.calculation_request.section.section_type == SectionType.ANGLE:
        lines.append(
            "Reducer для angle трактуется как инженерная композиция двух полок с вычетом зоны перекрытия. "
            "Он пригоден для остаточной оценки, но не является расчетом тонкостенной крутильной работы уголка."
        )

    if response.resistance_mode == ResistanceMode.APPROXIMATE:
        lines.append(
            "Несущая способность определялась в приближенном режиме. "
            "Для данного случая задействована укрупненная проверка сжатия через коэффициент устойчивости."
        )
    elif response.resistance_mode == ResistanceMode.COMPRESSION_ENHANCED:
        lines.append(
            "Несущая способность определялась в режиме compression_enhanced. "
            "Использована инженерная slenderness-редукция с явными входами по приведенной длине и закреплению, но не полный нормативный расчет по СП 16."
        )
    elif response.resistance_mode == ResistanceMode.COMBINED_BASIC:
        lines.append(
            "Несущая способность определялась в режиме combined_basic. "
            "Использовано базовое интеракционное соотношение N/Nrd + M/Mrd <= 1.0, а не полный нормативный расчет по СП 16."
        )
    elif response.resistance_mode == ResistanceMode.COMBINED_ENHANCED:
        lines.append(
            "Несущая способность определялась в режиме combined_enhanced. "
            "Использована инженерно-аппроксимированная комбинированная проверка с учетом продольной силы, изгиба и поправок второго порядка, "
            "но не полный нормативный расчет по СП 16."
        )

    if response.rate_fit_mode == RateFitMode.BASELINE_FALLBACK:
        lines.append(
            "История обследований недостаточна для идентификации наблюдаемой скорости деградации; "
            "прогноз продолжен по baseline-модели атмосферной коррозии."
        )
    elif response.rate_fit_mode == RateFitMode.SINGLE_OBSERVATION:
        lines.append(
            "Скорость деградации оценена по одному обследованию. "
            "Такая оценка пригодна как инженерный индикатор, но имеет повышенную неопределенность."
        )
    elif response.rate_fit_mode == RateFitMode.TWO_POINT:
        lines.append(
            "Скорость деградации оценена по двум точкам истории обследований. "
            "Тренд является приближенным и чувствителен к качеству замеров."
        )
    elif response.rate_fit_mode == RateFitMode.ROBUST_HISTORY_FIT_LOW_CONFIDENCE:
        lines.append(
            "Робастная аппроксимация истории обследований выполнена в режиме пониженной достоверности. "
            "Результат пригоден как инженерный ориентир, но требует осторожной интерпретации."
        )

    if response.ml_mode == "heuristic":
        lines.append(
            "Гибридный ML-контур отработал в эвристическом режиме. "
            "Он используется только для корректировки скорости деградации и не заменяет инженерный расчет сопротивления."
        )
    elif response.ml_mode == "fallback":
        lines.append(
            "Табличные ML-кандидаты недоступны или не обучены; использован детерминированный fallback-контур."
        )

    if response.engineering_confidence_level in (EngineeringConfidenceLevel.C, EngineeringConfidenceLevel.D):
        lines.append(
            f"Итоговый класс инженерной уверенности: {translate_confidence_level(response.engineering_confidence_level)}."
        )

    if response.risk_mode == "engineering_uncertainty_band" or getattr(response.risk_mode, "value", "") == "engineering_uncertainty_band":
        lines.append(
            "Показанный uncertainty band является инженерным интервалом guidance по сценариям и скорости деградации, "
            "а не формальной вероятностной оценкой надежности."
        )
        lines.append(
            f"Источник uncertainty: {translate_uncertainty_source(response.uncertainty_source)}; "
            f"уровень uncertainty: {translate_uncertainty_level(response.uncertainty_level)}."
        )

    if response.refinement_diagnostics.status in {"near_flat_no_crossing", "numerically_uncertain_crossing"}:
        lines.append(
            f"Диагностика поиска предельного состояния: {translate_refinement_status(response.refinement_diagnostics.status)}."
        )

    lines.append(response.risk_profile.method_note)
    lines.extend([f"Предупреждение: {warning}" for warning in response.warnings])
    lines.extend([f"Флаг fallback: {translate_fallback_flag(flag)}" for flag in response.fallback_flags])

    if not lines:
        lines.append(
            "Существенные fallback-режимы и укрупняющие допущения для текущего расчета не зафиксированы; "
            "расчет выполнен в прямом инженерном режиме в рамках реализованной модели."
        )

    return unique_lines(lines)


def sample_timeline(timeline: Sequence, max_rows: int) -> Sequence:
    if len(timeline) <= max_rows:
        return timeline

    step = max(1, len(timeline) // (max_rows - 1))
    sampled = list(timeline[::step])
    if sampled[-1] is not timeline[-1]:
        sampled.append(timeline[-1])
    return sampled[:max_rows]


def add_docx_paragraph(document: Document, text: str) -> None:
    document.add_paragraph(text)


def add_docx_table(document: Document, headers: Sequence[str], rows: Sequence[Sequence[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)

    document.add_paragraph("")


def register_pdf_fonts() -> Tuple[str, str]:
    regular_candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
    ]
    bold_candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf"),
    ]

    regular_path = first_existing_path(regular_candidates)
    bold_path = first_existing_path(bold_candidates)

    if regular_path and "CorrosionReportRegular" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("CorrosionReportRegular", str(regular_path)))
    if bold_path and "CorrosionReportBold" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("CorrosionReportBold", str(bold_path)))

    regular_font = "CorrosionReportRegular" if regular_path else "Helvetica"
    bold_font = "CorrosionReportBold" if bold_path else "Helvetica-Bold"
    return regular_font, bold_font


def first_existing_path(paths: Iterable[Path]) -> Optional[Path]:
    for path in paths:
        if path.exists():
            return path
    return None


def get_pdf_styles(regular_font: str, bold_font: str) -> dict:
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "ReportTitle",
            parent=styles["Title"],
            fontName=bold_font,
            fontSize=16,
            leading=20,
            spaceAfter=8,
        ),
        "heading": ParagraphStyle(
            "ReportHeading",
            parent=styles["Heading2"],
            fontName=bold_font,
            fontSize=11,
            leading=14,
            spaceBefore=6,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "ReportBody",
            parent=styles["BodyText"],
            fontName=regular_font,
            fontSize=9,
            leading=12,
            spaceAfter=4,
        ),
        "table": ParagraphStyle(
            "ReportTable",
            parent=styles["BodyText"],
            fontName=regular_font,
            fontSize=8,
            leading=10,
        ),
    }


def build_pdf_section(title: str, headers: Sequence[str], rows: Sequence[Sequence[str]], styles: dict) -> List:
    data = [list(headers)]
    for row in rows:
        data.append([Paragraph(escape(str(value)), styles["table"]) for value in row])

    table = Table(data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E2F3")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("FONTNAME", (0, 0), (-1, 0), styles["heading"].fontName),
                ("FONTNAME", (0, 1), (-1, -1), styles["table"].fontName),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#808080")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return [Paragraph(title, styles["heading"]), table, Spacer(1, 4 * mm)]


def build_pdf_list_section(title: str, lines: Sequence[str], styles: dict) -> List:
    content: List = [Paragraph(title, styles["heading"])]
    for line in lines:
        content.append(Paragraph(escape(f"• {line}"), styles["body"]))
    content.append(Spacer(1, 4 * mm))
    return content


def render_markdown_table(headers: Sequence[str], rows: Sequence[Sequence[str]]) -> str:
    header_line = "| " + " | ".join(escape_markdown_cell(value) for value in headers) + " |"
    separator_line = "| " + " | ".join(["---"] * len(headers)) + " |"
    body_lines = [
        "| " + " | ".join(escape_markdown_cell(value) for value in row) + " |"
        for row in rows
    ]
    return "\n".join([header_line, separator_line, *body_lines])


def render_html_table(headers: Sequence[str], rows: Sequence[Sequence[str]]) -> str:
    parts = ["    <table>", "      <thead>", "        <tr>"]
    parts.extend([f"          <th>{escape(str(header))}</th>" for header in headers])
    parts.extend(["        </tr>", "      </thead>", "      <tbody>"])
    for row in rows:
        parts.append("        <tr>")
        parts.extend([f"          <td>{escape(str(value))}</td>" for value in row])
        parts.append("        </tr>")
    parts.extend(["      </tbody>", "    </table>"])
    return "\n".join(parts)


def escape_markdown_cell(value: object) -> str:
    return str(value).replace("|", "\\|").replace("\n", " ")


def format_datetime(value: datetime) -> str:
    return value.astimezone(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")


def format_number(value: float, digits: int) -> str:
    return f"{value:.{digits}f}"


def format_optional_number(value: Optional[float], digits: int) -> str:
    if value is None:
        return "-"
    return format_number(value, digits)


def format_life_interval(interval) -> str:
    if interval is None:
        return "-"
    lower = format_optional_number(getattr(interval, "lower_years", None), 2)
    upper = format_optional_number(getattr(interval, "upper_years", None), 2)
    if lower == "-" and upper == "-":
        return "-"
    return f"[{lower}; {upper}]"


def build_action_rows(request: CalculationRequest) -> List[List[str]]:
    rows = [["Тип проверки", translate_check_type(request.action.check_type)]]
    if request.action.check_type in (
        CheckType.COMBINED_AXIAL_BENDING_BASIC,
        CheckType.COMBINED_AXIAL_BENDING_ENHANCED,
    ):
        rows.extend(
            [
                ["Тип продольной силы", translate_axial_force_kind(request.action.axial_force_kind)],
                ["Продольная сила", format_optional_number(request.action.axial_force_value, 3)],
                ["Изгибающий момент", format_optional_number(request.action.bending_moment_value, 3)],
            ]
        )
        if request.action.check_type == CheckType.COMBINED_AXIAL_BENDING_ENHANCED:
            rows.extend(
                [
                    ["Р­С„С„РµРєС‚РёРІРЅР°СЏ РґР»РёРЅР°", format_optional_number(request.action.effective_length_mm, 1)],
                    ["РљРѕСЌС„С„РёС†РёРµРЅС‚ РїСЂРёРІРµРґРµРЅРЅРѕР№ РґР»РёРЅС‹", format_optional_number(request.action.effective_length_factor, 3)],
                    ["РЈСЃР»РѕРІРёРµ Р·Р°РєСЂРµРїР»РµРЅРёСЏ", request.action.support_condition or "-"],
                    ["РљРѕСЌС„С„РёС†РёРµРЅС‚ СѓСЃРёР»РµРЅРёСЏ РјРѕРјРµРЅС‚Р°", format_optional_number(request.action.moment_amplification_factor, 3)],
                ]
            )
    else:
        rows.append(["Расчетное воздействие", format_optional_number(request.action.demand_value, 3)])
    rows.append(["Рост воздействия в год", format_number(request.action.demand_growth_factor_per_year, 4)])
    return rows


def build_action_rows(request: CalculationRequest) -> List[List[str]]:
    rows = [["Тип проверки", translate_check_type(request.action.check_type)]]
    if request.action.check_type in (
        CheckType.COMBINED_AXIAL_BENDING_BASIC,
        CheckType.COMBINED_AXIAL_BENDING_ENHANCED,
    ):
        rows.extend(
            [
                ["Тип продольной силы", translate_axial_force_kind(request.action.axial_force_kind)],
                ["Продольная сила", format_optional_number(request.action.axial_force_value, 3)],
                ["Изгибающий момент", format_optional_number(request.action.bending_moment_value, 3)],
            ]
        )
        if request.action.check_type == CheckType.COMBINED_AXIAL_BENDING_ENHANCED:
            rows.extend(
                [
                    ["Эффективная длина", format_optional_number(request.action.effective_length_mm, 1)],
                    ["Коэффициент приведенной длины", format_optional_number(request.action.effective_length_factor, 3)],
                    ["Условие закрепления", request.action.support_condition or "-"],
                    ["Коэффициент усиления момента", format_optional_number(request.action.moment_amplification_factor, 3)],
                ]
            )
    elif request.action.check_type == CheckType.AXIAL_COMPRESSION_ENHANCED:
        rows.extend(
            [
                ["Расчетное воздействие", format_optional_number(request.action.demand_value, 3)],
                ["Эффективная длина", format_optional_number(request.action.effective_length_mm, 1)],
                ["Коэффициент приведенной длины", format_optional_number(request.action.effective_length_factor, 3)],
                ["Условие закрепления", request.action.support_condition or "-"],
            ]
        )
    else:
        rows.append(["Расчетное воздействие", format_optional_number(request.action.demand_value, 3)])
    rows.append(["Рост воздействия в год", format_number(request.action.demand_growth_factor_per_year, 4)])
    return rows


def translate_check_type(check_type: CheckType) -> str:
    mapping = {
        CheckType.AXIAL_TENSION: "axial_tension - растяжение",
        CheckType.AXIAL_COMPRESSION: "axial_compression - сжатие",
        CheckType.AXIAL_COMPRESSION_ENHANCED: "axial_compression_enhanced - сжатие (engineering enhanced)",
        CheckType.BENDING_MAJOR: "bending_major - изгиб по главной оси",
        CheckType.COMBINED_AXIAL_BENDING_BASIC: "combined_axial_bending_basic - базовая комбинированная проверка",
    }
    mapping[CheckType.COMBINED_AXIAL_BENDING_ENHANCED] = "combined_axial_bending_enhanced - усиленная комбинированная проверка"
    return mapping[check_type]


def translate_axial_force_kind(kind: AxialForceKind) -> str:
    mapping = {
        AxialForceKind.TENSION: "tension - растяжение",
        AxialForceKind.COMPRESSION: "compression - сжатие",
    }
    return mapping[kind]


def translate_forecast_mode(mode: ForecastMode) -> str:
    mapping = {
        ForecastMode.BASELINE: "baseline - классическая атмосферная модель",
        ForecastMode.OBSERVED: "observed - прогноз по наблюдаемой скорости",
        ForecastMode.HYBRID: "hybrid - наблюдения + коррекция baseline",
    }
    return mapping[mode]


def translate_confidence_level(level: EngineeringConfidenceLevel) -> str:
    mapping = {
        EngineeringConfidenceLevel.A: "A - прямой reducer и прямой расчет сопротивления",
        EngineeringConfidenceLevel.B: "B - прямой reducer и приближенная инженерная проверка",
        EngineeringConfidenceLevel.C: "C - fallback-допущения по сечению или режиму расчета",
        EngineeringConfidenceLevel.D: "D - неполные данные или исследовательский режим",
    }
    return mapping[level]


def translate_resistance_mode(mode: ResistanceMode | str) -> str:
    normalized = mode.value if isinstance(mode, ResistanceMode) else str(mode)
    mapping = {
        ResistanceMode.DIRECT.value: "direct - прямая инженерная проверка",
        ResistanceMode.APPROXIMATE.value: "approximate - приближенная инженерная проверка",
        ResistanceMode.COMPRESSION_ENHANCED.value: "compression_enhanced - усиленное инженерное сжатие",
        ResistanceMode.COMBINED_BASIC.value: "combined_basic - базовая комбинированная проверка",
    }
    mapping[ResistanceMode.COMBINED_ENHANCED.value] = "combined_enhanced - усиленная комбинированная проверка"
    return mapping.get(normalized, normalized)


def translate_reducer_mode(mode: ReducerMode | str) -> str:
    normalized = mode.value if isinstance(mode, ReducerMode) else str(mode)
    mapping = {
        ReducerMode.DIRECT.value: "direct - прямой редьюсер профиля",
        ReducerMode.GENERIC_FALLBACK.value: "generic_fallback - только fallback через generic_reduced",
    }
    return mapping.get(normalized, normalized)


def translate_rate_fit_mode(mode: RateFitMode | str) -> str:
    normalized = mode.value if isinstance(mode, RateFitMode) else str(mode)
    mapping = {
        RateFitMode.BASELINE_FALLBACK.value: "baseline_fallback - без истории обследований",
        RateFitMode.SINGLE_OBSERVATION.value: "single_observation - одно обследование",
        RateFitMode.TWO_POINT.value: "two_point - две точки истории",
        RateFitMode.ROBUST_HISTORY_FIT.value: "robust_history_fit - робастная аппроксимация истории 3+ точек",
        RateFitMode.ROBUST_HISTORY_FIT_LOW_CONFIDENCE.value: "robust_history_fit_low_confidence - робастная аппроксимация с пониженной уверенностью",
    }
    return mapping.get(normalized, normalized)


def translate_risk_mode(mode) -> str:
    normalized = mode.value if hasattr(mode, "value") else str(mode)
    if normalized == "engineering_uncertainty_band":
        return "engineering_uncertainty_band - инженерный интервал неопределенности"
    if normalized == "scenario_risk":
        return "scenario_risk - сценарный индикатор риска"
    return normalized


def translate_uncertainty_level(level) -> str:
    normalized = level.value if hasattr(level, "value") else str(level)
    mapping = {
        "low": "low - узкий инженерный интервал",
        "moderate": "moderate - умеренный инженерный разброс",
        "high": "high - повышенная инженерная неопределенность",
        "very_high": "very_high - очень высокая инженерная неопределенность",
    }
    return mapping.get(normalized, normalized)


def translate_uncertainty_source(source: str) -> str:
    mapping = {
        "scenario_library_only": "scenario_library_only - только базовая сценарная библиотека",
        "inspection_history_band": "inspection_history_band - полоса по данным обследований",
        "inspection_history_limited": "inspection_history_limited - история обследований ограничена",
        "inspection_history_with_baseline_fallback": "inspection_history_with_baseline_fallback - часть зон добрана baseline fallback",
    }
    return mapping.get(str(source), str(source))


def translate_crossing_mode(mode: str) -> str:
    mapping = {
        "no_timeline": "нет временной диаграммы",
        "already_reached": "предельное состояние уже достигнуто",
        "coarse_bracket_linear": "линейная интерполяция в грубом интервале",
        "coarse_bracket_refined": "уточнение внутри грубого интервала",
        "no_crossing_within_horizon": "в пределах горизонта переход не найден",
        "coarse_only": "только грубая оценка",
    }
    return mapping.get(str(mode), str(mode))


def translate_refinement_status(status: str) -> str:
    mapping = {
        "no_timeline": "нет временной диаграммы",
        "already_reached": "предельное состояние уже достигнуто",
        "bracketed_crossing": "переход найден в устойчивом интервале",
        "numerically_uncertain_crossing": "переход найден, но интервал численно неустойчив",
        "near_flat_no_crossing": "в горизонте нет перехода, но кривая почти плоская",
        "no_crossing_within_horizon": "в горизонте переход не найден",
        "coarse_only": "только грубая оценка",
    }
    return mapping.get(str(status), str(status))


def translate_ml_mode(mode: str) -> str:
    mapping = {
        "heuristic": "heuristic - эвристическая корректировка скорости",
        "trained": "trained - обученный ансамбль кандидатов с резервным якорем",
        "fallback": "fallback - детерминированный резерв без табличных кандидатов",
    }
    return mapping.get(str(mode), str(mode))


def translate_fallback_flag(flag: str) -> str:
    if flag == "generic_reduced":
        return "эффективное сечение получено через generic_reduced как явный fallback-режим."
    if flag.startswith("forecast_source:") and flag.endswith(":baseline"):
        _, zone_id, _ = flag.split(":", 2)
        return f"зона {zone_id}: прогноз продолжен по baseline-модели."
    if flag.startswith("resistance_mode:"):
        return f"режим сопротивления {translate_resistance_mode(flag.split(':', 1)[1])}."
    return flag


def build_trajectory_summary(trajectories) -> str:
    central = len(getattr(trajectories, "central", []) or [])
    conservative = len(getattr(trajectories, "conservative", []) or [])
    upper = len(getattr(trajectories, "upper", []) or [])
    if not any([central, conservative, upper]):
        return "-"
    return f"central={central}, conservative={conservative}, upper={upper}"


def format_acceptance_policy(policy: dict) -> str:
    if not policy:
        return "-"
    return ", ".join(f"{key}={value}" for key, value in sorted(policy.items()))


def summarize_dataset_journal(dataset_journal: Sequence[dict]) -> str:
    items = []
    for item in dataset_journal:
        items.append(
            f"{item.get('dataset_kind', '?')}:{item.get('version', '?')}:{item.get('dataset_hash', '-')}"
        )
    return "; ".join(items) or "-"


def summarize_candidate_registry(candidate_registry: Sequence[dict]) -> str:
    items = []
    for item in candidate_registry:
        items.append(
            f"{item.get('name', '?')}={item.get('status', '?')}:{item.get('reason', '-')}"
        )
    return "; ".join(items) or "-"


def unique_lines(lines: Sequence[str]) -> List[str]:
    seen = set()
    unique: List[str] = []
    for line in lines:
        if line in seen:
            continue
        seen.add(line)
        unique.append(line)
    return unique
